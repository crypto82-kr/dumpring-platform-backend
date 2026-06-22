import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        borders.append(border)
    tblPr.append(borders)

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("덤프링 (Dumpring) 코드/데이터베이스 시스템 용어 검토표")
    run_title.font.name = 'Malgun Gothic'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    
    doc.add_paragraph("본 문서는 덤프링(Dumpring) 실제 데이터베이스 스키마 및 소스코드 상에 변수, 테이블, 상태 코드로 사용 중인 실제 시스템 용어만을 추출한 목록입니다. 시스템 필드명 변경안을 기재하여 코드 리팩토링 및 표준화에 적용해 주십시오.\n")
    
    # Tables, Fields, Statuses used in code
    system_terms = [
        # --- [1] 데이터베이스 테이블명 ---
        ("users (Table)", "통합 사용자 계정 관리 테이블 (CI, 번호, 다중 역할 등)"),
        ("construction_sites (Table)", "공사현장 기본정보 및 현장 위경도, 지오펜싱 반경 관리"),
        ("site_employees (Table)", "현장에 소속된 직원 선등록(user_id=Null) 및 실가입 매칭 관계"),
        ("unloading_sites (Table)", "지주가 등록한 사토장 마스터 및 수용 가능 토질 콤마 스트링"),
        ("cars (Table)", "차량번호 및 톤수, 소유 차주 ID가 매핑된 덤프트럭 차량 테이블"),
        ("drivers (Table)", "기사의 선등록 번호(registered_phone) 및 임시 배정된 차량 매핑"),
        ("user_uploaded_documents (Table)", "가입 심사 서류 저장 및 파일명 연동 정보"),
        ("orders (Table - Legacy)", "구버전 오더 마스터 테이블 (상/하차지 위치 및 상태)"),
        ("site_profiles (Table)", "현장관리자/담당자의 회사명, 사업자번호 등 세부 프로필"),
        ("drop_off_profiles (Table)", "하차지주의 허가번호, 사토장 주소 세부 프로필"),
        ("site_user_mappings (Table)", "현장담당자와 현장 간의 다대다 매핑 신청 및 소장 심사 관계"),
        ("common_codes (Table)", "토사종류, 트럭규격, 정산방식 등 동적 코드 관리 테이블"),
        ("drop_offs (Table)", "지주가 개설한 사토지 마스터 데이터 테이블"),
        ("drop_off_requests (Table)", "사토지 지주가 등록하는 목표 수량(Tonnage) 수용 공고 테이블"),
        ("job_posts (Table)", "상하차지 B2B 매칭 및 배차 모집 통합 오더 테이블"),
        ("driver_favorite_regions (Table)", "기사가 등록해 두는 선호 배차 시도/시군구 구역 정보"),
        ("dispatch_tickets (Table)", "기사 운행 GPS 트래킹 및 실시간 정산 처리 기록"),
        ("sdui_themes (Table)", "앱의 색상, 폰트 테마 서버 제어용 데이터"),
        ("wallets (Table)", "기사/차주/지주별 가상 지갑 포인트 잔고 관리 테이블"),
        ("wallet_transactions (Table)", "지갑 내 포인트 이체, 충금 등 상세 거래 기록"),
        ("gps_tracks (Table)", "주행 티켓별 실시간 위경도 좌표 및 증분 운임 수집 데이터"),
        
        # --- [2] 핵심 테이블 컬럼 / 변수명 ---
        ("ci", "본인인증 시 수신하는 개인 고유 해시 키 (중복가입 차단용)"),
        ("phone_number", "사용자 계정의 로그인 아이디 역할을 하는 휴대폰 번호"),
        ("is_site_manager", "유저의 공사현장 소장/관리자 권한 토글 필드 (Boolean)"),
        ("is_site_worker", "유저의 현장 진출입 통제 담당자 권한 토글 필드 (Boolean)"),
        ("is_owner", "유저의 덤프 소유 차주 권한 토글 필드 (Boolean)"),
        ("is_driver", "유저의 덤프 트럭 운전기사 권한 토글 필드 (Boolean)"),
        ("is_drop_off", "유저의 사토장 지주 권한 토글 필드 (Boolean)"),
        ("is_approved", "플랫폼 어드민의 회원가입 승인 처리 필드 (Boolean)"),
        ("reject_reason", "가입 승인 거절 시 입력되는 사유 필드"),
        ("site_key", "현장에 소속 직원들이 가입 연동할 수 있게 발급하는 고유 식별키"),
        ("geofencing_radius", "하차지 도착 자동 감지를 위해 설정하는 원형 GPS 경계 반경 (m)"),
        ("preferred_soil_types", "사토장이 허용하는 토사 종류 콤마 연결 문자열"),
        ("registered_phone", "차주/소장이 기사/직원을 선등록할 때 입력하는 대상 번호"),
        ("current_car_id", "기사에게 현재 배정되어 운행에 사용되는 차량 외래키 ID"),
        ("document_code", "필수 서류 구분을 위한 공통코드 연동 코드"),
        ("soil_type / material_type", "오더 발행 시 지정한 토사 종류 (GOOD_SOIL, MUD_SOIL 등)"),
        ("required_tonnage / truck_type", "오더 발행 시 필요한 덤프 규격 톤수"),
        ("payer_type", "운반비/덤프비 정산 시 지불 주체 코드"),
        ("payment_method", "운송 완료 후 정산 결제 방식 (MONTHLY: 월대, DAILY: 당일)"),
        ("unit_price / offered_unit_price", "하차지 수용 단가 및 상차지 제시 정산 단가"),
        ("accumulated_fare", "실시간 GPS 운행 시 병산 가산되어 누적되는 정산 운임"),
        ("drive_distance_km", "주행 시작 시점부터 실시간 수집되어 누적된 총 거리 (km)"),
        ("drive_time_seconds", "주행 중 신호대기 및 운행에 경과한 누적 시간 (초)"),
        
        # --- [3] 상태 및 분류값 (Enum / Code) ---
        ("WAITING_MATCH", "JobPost 상태 - 상차지 우선 등록 후 하차지 매칭 대기 중"),
        ("WAITING_APPROVAL", "JobPost 상태 - 매칭 대상이 지정되어 상차지/하차지 승인 대기 중"),
        ("OPEN", "JobPost 상태 - 매칭 최종 승인 완료되어 기사들에게 모집 공고 노출 중"),
        ("COMPLETED", "JobPost/Ticket 상태 - 하차 완료 승인되어 운송 및 정산 최종 종결"),
        ("CANCELLED", "JobPost/Ticket 상태 - 현장 취소 또는 거절 등으로 오더 파기"),
        ("ACCEPTED", "Ticket 상태 - 기사가 배차 모집을 선택해 수락 완료한 운행 대기 상태"),
        ("DRIVING", "Ticket 상태 - 기사가 운행시작을 눌러 GPS 미터기가 실시간 가동 중인 주행 상태"),
        ("ARRIVED / ARRIVED_DROPZONE", "Ticket 상태 - 하차지 반경 지오펜싱(200m) 영역 진입 자동 감지 상태"),
        ("APPROVED", "Ticket 상태 - 하차지 지주가 토질 검사 통과 및 정상 반입 승인 완료한 정산 확정 상태"),
        ("REJECTED", "Ticket 상태 - 토질 부적합 판정 등으로 반입을 보류 및 거부한 상태"),
        ("DISPUTED", "Ticket 상태 - 반입 거부 또는 요금 측정 오류로 정산 분쟁이 접수된 상태"),
        ("SETTLE_ADJUSTED", "Ticket 상태 - 분쟁 접수 후 어드민이 직권으로 70% 정산 등을 실행 완료한 상태"),
        ("PENDING", "SiteUserMapping 상태 - 담당자의 현장 소속 신청 후 소장 승인을 대기하는 상태")
    ]
    
    # Table creation
    table = doc.add_table(rows=1, cols=3)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '시스템/코드 내 명칭'
    hdr_cells[1].text = '바꿀 명칭 (변경 필드명)'
    hdr_cells[2].text = '코드 상의 역할 및 기능 설명'
    
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, 'D97706') # Amber Orange Header
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    col_widths = [Inches(2.5), Inches(2.2), Inches(3.3)]
            
    for idx, (curr_name, desc) in enumerate(system_terms):
        row = table.add_row()
        cells = row.cells
        cells[0].text = curr_name
        cells[1].text = "" # Blank for custom modification
        cells[2].text = desc
        
        bg_color = 'F9F8F6' if idx % 2 == 0 else 'FFFFFF'
        for i, cell in enumerate(cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                else:
                    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                    
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Save
    file_path = "D:\\Projects\\dumpring\\dumpring-platform-backend\\docs\\덤프링_시스템_코드_용어_정의표.docx"
    doc.save(file_path)
    print(f"System Glossary Word document successfully created at: {file_path}")

if __name__ == "__main__":
    main()

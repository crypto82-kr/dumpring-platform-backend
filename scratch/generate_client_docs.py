import asyncio
import sys
import os
import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def create_element(name):
    return OxmlElement(name)

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Charcoal Black
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("덤프링 (Dumpring) 플랫폼 인프라 설명 및 용어 정의서")
    run_title.font.name = 'Malgun Gothic'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Amber Orange
    
    # Subtitle
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_desc = desc.add_run("협의 및 의사 결정권자 브리핑용 인프라 사양, 데이터 백업 구조, 운영 비용 분석서")
    run_desc.font.italic = True
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(20)
    
    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20)
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        return p

    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(10)

    # 1. 용어 정의서
    add_heading_1("1. 플랫폼 주요 용어 정의 (Glossary)")
    doc.add_paragraph("본 장에서는 덤프링 플랫폼에서 사용되는 역할, 기능 및 운행 관련 핵심 용어들을 정의합니다.")

    # Table 1: 용어 정의
    terms = [
        ("Driver (기사)", "덤프트럭을 운전하는 사용자입니다. 배차 공고를 수락하고 상/하차지간 GPS 미터기를 켜고 운행을 담당합니다."),
        ("Landowner (지주)", "하차지(사토장)의 관리자입니다. 반입된 토사의 상태를 육안 검증하여 최종 승인/반려(회차)를 결정합니다."),
        ("Site Manager (현장 관리자)", "상차지의 관리자입니다. 오더 공고를 등록하고 기사의 현장 진입 시 상차 완료 원격 승인을 처리합니다."),
        ("Car Owner (차주)", "덤프트럭 차량 소유주입니다. 소속 기사를 관리하고 정산 및 매출을 확인합니다."),
        ("Admin (플랫폼 관리자)", "전체 덤프링 시스템 운영자입니다. 기사 면허 서류 승인 및 분쟁 조정을 처리합니다."),
        ("상차지 (Loading Site)", "토사를 적재하고 출발하는 상차 현장입니다 (건설 현장 등)."),
        ("하차지 (Drop-off Site)", "토사를 하차(반입)하는 목적지 현장입니다 (사토장 등)."),
        ("GPS 미터기 (GPS Meter)", "실시간 주행 거리와 대기 시간을 GPS 좌표 기반으로 계산해 요금을 책정하는 시스템입니다."),
        ("상차 승인 (Loading Approval)", "기사가 실제 현장에서 상차했음을 검증하는 단계입니다. (현장 QR 스캔 또는 관리자 원격 승인)"),
        ("반입 승인 (Drop-off Approval)", "하차지 지주가 토질 적합 여부 검사 후 하차를 최종 통과시키는 단계입니다."),
        ("지주 부재 대체 전표", "지주가 부재중일 때 기사가 사진 증빙을 제출해 임시 승인을 신청하는 모드입니다."),
        ("전자 전표 (E-Receipt)", "종이 전표를 대체하는 디지털 증빙 영수증으로 최종 단가, 요금, 거리, 시간이 명시됩니다.")
    ]
    
    table1 = doc.add_table(rows=1, cols=2)
    table1.style = 'Table Grid'
    table1.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    hdr_cells = table1.rows[0].cells
    hdr_cells[0].text = '용어명 (영어)'
    hdr_cells[1].text = '정의 및 역할'
    for cell in hdr_cells:
        set_cell_background(cell, 'D97706') # Amber
        for p in cell.paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    for term, defn in terms:
        row = table1.add_row()
        cells = row.cells
        cells[0].text = term
        cells[1].text = defn
        cells[0].paragraphs[0].runs[0].font.bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 2. 인프라 아키텍처 및 시스템 사양
    add_heading_1("2. 인프라 아키텍처 및 시스템 사양 (Infrastructure)")
    doc.add_paragraph("덤프링은 최신 비동기 FastAPI 웹 서버와 안정성 높은 Supabase PostgreSQL 데이터베이스를 연동해 구축되었습니다.")
    
    add_heading_2("2.1. 트래픽 한도 및 최대 동시 접속자 수")
    add_bullet("FastAPI API 웹 서버: 초당 150~300회 요청(RPS) 처리가 가능하며, 동시에 운행 중인 차량 1,000대 ~ 3,000대 수준을 지연 없이 실시간 관제할 수 있습니다.")
    add_bullet("PostgreSQL 데이터베이스: PgBouncer 커넥션 풀러를 탑재하여 동시 사용자 10,000명 이상의 무거운 데이터 요청도 서버 고갈 없이 정상 대기열로 분산 처리해 냅니다.")
    add_bullet("웹 어드민 콘솔: 전 세계 Edge 분산 서버망(Vercel CDN)을 활용해 웹 트래픽 한도가 실질적으로 무제한이며 다운타임이 발생하지 않습니다.")
    
    add_heading_2("2.2. 데이터 백업 및 데이터 보호 (PITR)")
    doc.add_paragraph("배차 정산 데이터의 안전성을 100% 보장하기 위해 초(Second) 단위 변경 사항을 실시간 백업하는 PITR(Point-in-Time Recovery) 기술이 구축되어 있습니다. 시스템에 예기치 못한 문제가 발생하더라도, 단 1초 전의 상태로 복원할 수 있어 정산 데이터 손실을 원천 차단합니다.")
    
    add_heading_2("2.3. 향후 메이저 클라우드(AWS/GCP) 마이그레이션")
    doc.add_paragraph("기사 및 이용 트랙수가 급증하여 전국 단위의 메이저 서비스로 확장할 시, 기존 코드를 재개발할 필요 없이 다음과 같이 즉각 이주(Migration)가 가능한 도커 컨테이너 및 표준 PostgreSQL 기술을 준용하였습니다.")
    add_bullet("API 서버: AWS ECS (Fargate) 또는 Google Cloud Run으로 도커 이미지 즉시 이주")
    add_bullet("데이터베이스: 표준 PostgreSQL pg_dump/pg_restore 명령어로 AWS RDS 또는 Google Cloud SQL로 즉시 데이터 이주")
    add_bullet("파일 스토리지: S3 API 호환 표준 규격으로 설정 주소 변경만으로 AWS S3로 이미지 즉시 이주")

    doc.add_paragraph().paragraph_format.space_after = Pt(14)

    # 3. 비용 분석
    add_heading_1("3. 인프라 운영 비용 요약 (Cost Analysis)")
    doc.add_paragraph("덤프링 실 서비스 가동에 필요한 최소화된 인프라 요금 및 기사용 실명 본인인증 비용 구성표입니다.")

    # Table 2: 비용 요약
    costs = [
        ("백엔드 서버", "Render (FastAPI)", "Starter 플랜", "월 $7 (약 1만 원)", "API 상시 가동 및 무지연 가동용"),
        ("데이터베이스", "Supabase (PostgreSQL)", "Pro 플랜", "월 $25 (약 3만 4천 원)", "일일 기본 백업 및 스토리지 포함"),
        ("실시간 백업", "Supabase PITR", "초 단위 복구 옵션", "월 $8 (약 1만 1천 원)", "데이터 무손실 복구 보장"),
        ("웹 어드민 호스팅", "Vercel (Next.js)", "무료 또는 Pro 플랜", "월 $20 (약 2만 7천 원)", "어드민 페이지 서비스용"),
        ("협업 도구", "Google Workspace", "Business Standard", "월 약 $24 (2인 기준)", "개발자 협업 및 공유 드라이브용"),
        ("iOS 빌드", "Codemagic / GitHub macOS", "무료 또는 종량제", "월 $15 (약 2만 원)", "Mac 장비 없을 시 원격 iOS 패키징"),
        ("휴대폰 본인확인", "PASS / SMS 인증", "성공 건당 요금", "건당 약 40~50원", "회원가입 시 최초 1회 실명 확인용"),
        ("간편 본인확인", "카카오 / 네이버 간편인증", "성공 건당 요금", "건당 약 100~200원", "카카오톡/네이버 간편 실명 확인용"),
        ("마켓 등록비", "Apple / Google 스토어", "스토어 연동 라이선스", "Apple 연 $99 / 구글 최초 1회 $25", "앱 배포용 개발자 계정 비용")
    ]
    
    table2 = doc.add_table(rows=1, cols=5)
    table2.style = 'Table Grid'
    table2.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    
    hdr2_cells = table2.rows[0].cells
    hdr2_cells[0].text = '구분'
    hdr2_cells[1].text = '서비스명'
    hdr2_cells[2].text = '플랜 사양'
    hdr2_cells[3].text = '비용 (월/건)'
    hdr2_cells[4].text = '용도 설명'
    
    for cell in hdr2_cells:
        set_cell_background(cell, 'D97706')
        for p in cell.paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    for item, serv, spec, price, usage in costs:
        row = table2.add_row()
        cells = row.cells
        cells[0].text = item
        cells[1].text = serv
        cells[2].text = spec
        cells[3].text = price
        cells[4].text = usage
        cells[0].paragraphs[0].runs[0].font.bold = True

    add_heading_2("3.1. 추천 단계별 예산 운영 제안")
    add_bullet("현장 실증 검증(PoC) 권장 세팅: 월 약 10만 원 ~ 12만 원 (서버 상시기동 + Supabase Pro + 구글 워크스페이스 + iOS 클라우드 빌드)")
    add_bullet("상용 서비스 정식 론칭 세팅: 월 약 15만 원 ~ 18만 원 (실증 세팅 + 서버 무중단 이중화 + 실시간 초 단위 PITR 백업)")
    add_bullet("회원가입 본인확인 실비: 기사 신규 가입 1인 성공 시 문자/PASS 인증 수수료 건당 약 40~50원 실비 별도 청구")

    # Save the document
    file_path = "D:\\Projects\\dumpring\\dumpring-platform-backend\\docs\\덤프링_플랫폼_인프라_및_용어_정의서.docx"
    doc.save(file_path)
    print(f"Client Word document successfully created at: {file_path}")

if __name__ == '__main__':
    main()

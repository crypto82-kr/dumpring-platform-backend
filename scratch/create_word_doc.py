import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, hex_color):
    """Set the background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def create_element(name):
    return OxmlElement(name)

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Malgun Gothic'
    font.size = Pt(10)
    font.color.rgb = RGBColor(0x1F, 0x29, 0x37) # Charcoal Black
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("덤프링 (Dumpring) 서버 기동 및 배포 매뉴얼")
    run_title.font.name = 'Malgun Gothic'
    run_title.font.size = Pt(20)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xD9, 0x77, 0x06) # Amber Orange
    
    # Subtitle or description
    desc = doc.add_paragraph()
    desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_desc = desc.add_run("플랫폼 백엔드 개발자 및 운영자를 위한 실무 실행 지침서")
    run_desc.font.italic = True
    run_desc.font.size = Pt(11)
    run_desc.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    doc.add_paragraph().paragraph_format.space_after = Pt(24)
    
    # Heading 1 helper
    def add_custom_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
        return p

    # Heading 2 helper
    def add_custom_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(12)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        return p
        
    # Code block helper
    def add_code_block(text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        cell = table.cell(0, 0)
        set_cell_background(cell, 'F3F4F6') # Light gray background
        
        # Border style
        tcPr = cell._tc.get_or_add_tcPr()
        borders = create_element('w:tcBorders')
        for border_name in ['top', 'left', 'bottom', 'right']:
            border = create_element(f'w:{border_name}')
            border.set(qn('w:val'), 'single')
            border.set(qn('w:sz'), '4')
            border.set(qn('w:space'), '0')
            border.set(qn('w:color'), 'E5E7EB')
            borders.append(border)
        tcPr.append(borders)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        
        run = p.add_run(text)
        run.font.name = 'Consolas'
        run.font.size = Pt(9.5)
        run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Content
    add_custom_heading_1("1. 로컬 개발 서버 기동 (Local Run)")
    
    p = doc.add_paragraph("본 장에서는 로컬 PC 환경에서 FastAPI 백엔드 서버를 설정하고 기동하는 가이드를 제공합니다.")
    
    add_custom_heading_2("1.1. 가상환경 활성화 및 패키지 설치")
    doc.add_paragraph("프로젝트 최상위 루트 디렉토리에서 터미널을 열고 가상환경을 활성화한 뒤 패키지를 설치합니다.")
    
    add_code_block(
        "# Windows (PowerShell)\n"
        ".\\venv\\Scripts\\Activate.ps1\n\n"
        "# Linux / macOS\n"
        "source venv/bin/activate\n\n"
        "# 종속성 패키지 설치\n"
        "pip install -r requirements.txt"
    )
    
    add_custom_heading_2("1.2. 환경 변수 설정 (.env)")
    doc.add_paragraph("루트 디렉토리에 .env 파일을 생성하고 아래와 같이 Supabase DB 연결 문자열 정보를 기입합니다.")
    
    add_code_block(
        "PROJECT_NAME=\"덤프링 (dumpring)\"\n"
        "SECRET_KEY=\"SUPER_SECRET_KEY_REPLACE_THIS_IN_PRODUCTION\"\n"
        "DATABASE_URL=\"postgresql+asyncpg://postgres.vvdydhxmwdrrazwyjvvu:crypto.co.kr@aws-1-ap-northeast-2.pooler.supabase.com:6543/postgres\""
    )
    
    p_warn = doc.add_paragraph()
    run_warn_tag = p_warn.add_run("[주의사항] ")
    run_warn_tag.font.bold = True
    run_warn_tag.font.color.rgb = RGBColor(0xEF, 0x44, 0x44)
    run_warn_txt = p_warn.add_run("Supabase 연결 시 커넥션 풀을 정상적으로 사용하고 커넥션 수 고갈을 막기 위해 포트는 반드시 6543 포트를 사용해야 합니다.")
    run_warn_txt.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
    
    add_custom_heading_2("1.3. 서버 실행")
    doc.add_paragraph("Uvicorn을 기동하여 로컬 루프백 인터페이스에서 서버를 실행합니다.")
    
    add_code_block("uvicorn app.main:app --reload --host 127.0.0.1 --port 8000")
    doc.add_paragraph("서버 실행 후 브라우저에서 아래 Swagger UI 주소로 접속해 정상 기동 여부를 확인합니다:")
    doc.add_paragraph("• 주소: http://127.0.0.1:8000/docs")
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_custom_heading_1("2. 데이터베이스 마이그레이션 및 초기 데이터 설정")
    doc.add_paragraph("가상환경이 활성화된 상태에서 아래 명령어를 실행하여 테이블 스키마 생성 및 초기 데이터를 적재합니다.")
    
    add_custom_heading_2("2.1. Alembic 마이그레이션 적용")
    doc.add_paragraph("테이블 수정 혹은 신규 생성 시 최신 버전 스키마를 Supabase DB에 마이그레이션합니다.")
    
    add_code_block(
        "# 최신 DB 스키마 업데이트\n"
        "alembic upgrade head\n\n"
        "# 만약 바로 이전 스키마 버전으로 롤백이 필요한 경우\n"
        "alembic downgrade -1"
    )
    
    add_custom_heading_2("2.2. 마스터 코드 및 테스트 데이터 로딩 (Seeding)")
    doc.add_paragraph("시스템 기동에 필수적인 토사 종류, 트럭 규격 등의 공통코드 마스터와 모의 시나리오 데이터를 입력합니다.")
    
    add_code_block(
        "# 공통코드 및 기초 마스터 코드 설정\n"
        "python seed_data.py\n\n"
        "# 테스트용 시나리오 데이터 설정\n"
        "python seed_scenario.py"
    )
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    add_custom_heading_1("3. 운영 서버 배포 (Render.com Deployment)")
    doc.add_paragraph("덤프링 백엔드 API 서버는 Render.com을 통해 운영 클라우드에 배포됩니다. GitHub main 브랜치와 연동되어 자동 배포(CD)가 수행됩니다.")
    
    add_custom_heading_2("3.1. Render 웹 서비스 설정값")
    doc.add_paragraph("Render 대시보드에서 New Web Service 생성 시 다음 사양 정보를 올바르게 기입합니다.")
    
    # Bullet points
    def add_bullet(text):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.font.name = 'Malgun Gothic'
        run.font.size = Pt(10)
    
    add_bullet("Name: dumpring-api")
    add_bullet("Region: Singapore (ap-southeast-1)")
    add_bullet("Branch: main")
    add_bullet("Runtime: Python")
    add_bullet("Build Command: pip install -r requirements.txt")
    add_bullet("Start Command: uvicorn app.main:app --host 0.0.0.0 --port $PORT")
    
    add_custom_heading_2("3.2. 환경 변수 (Environment Variables) 설정")
    doc.add_paragraph("Render 서비스 설정 내 Environment 탭에서 다음의 변수명과 값을 추가합니다.")
    
    add_bullet("DATABASE_URL: postgresql+asyncpg://postgres.vvdydhxmwdrrazwyjvvu:crypto.co.kr@aws-1-ap-northeast-2.pooler.supabase.com:6543/postgres")
    add_bullet("PYTHON_VERSION: 3.11.0")
    add_bullet("SECRET_KEY: [운영 환경 전용 암호화 키]")
    
    add_custom_heading_2("3.3. 수동 즉시 배포 트리거")
    doc.add_paragraph("코드를 푸시하지 않고 Render Dashboard 상에서 강제 재빌드 및 배포를 갱신하려면:")
    doc.add_paragraph("1. Render 대시보드 해당 Web Service의 오른쪽 상단에 위치한 Manual Deploy 드롭다운 메뉴를 클릭합니다.")
    doc.add_paragraph("2. Clear build cache & deploy 버튼을 클릭하여 수동 배포를 진행합니다.")
    
    # Save the document
    file_path = "D:\\Projects\\dumpring\\dumpring-platform-backend\\docs\\덤프링_서버_운영_및_배포_매뉴얼.docx"
    doc.save(file_path)
    print(f"Word document successfully created at: {file_path}")

if __name__ == "__main__":
    main()

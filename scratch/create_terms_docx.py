import docx
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH

def set_cell_background(cell, hex_color):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tc_pr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for margin, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{margin}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table):
    tblPr = table._tbl.tblPr
    borders = OxmlElement('w:tblBorders')
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), 'CCCCCC')
        borders.append(border)
    tblPr.append(borders)

def main():
    doc = Document()
    
    # Page setup
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    
    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title.add_run("덤프링 (Dumpring) 핵심 용어 정의 및 변경 검토표")
    run_title.font.name = 'Malgun Gothic'
    run_title.font.size = Pt(16)
    run_title.font.bold = True
    run_title.font.color.rgb = RGBColor(0xD9, 0x77, 0x06)
    
    doc.add_paragraph("본 문서는 덤프링(Dumpring) 서비스에 직접 연관되는 핵심 비즈니스 용어 변경 검토를 위한 양식입니다. '바꿀 이름' 열에 의견을 적어 표준화에 사용해 주시기 바랍니다.\n")
    
    terms = [
        ("상차지 (Loading Site)", "흙을 싣고 출발하는 공사 현장"),
        ("하차지 (Unloading Site / Drop-Off)", "실어 온 흙을 부리는 도착 사토장/매립지"),
        ("사토장 (Drop-Off Site)", "건설 잔토 및 사토를 합법적으로 묻거나 쌓는 매립 부지"),
        ("토사 / 사토 (Material / Spoil)", "건설 현장에서 발생해 운송 및 매립해야 하는 흙과 돌"),
        ("토종 (Material Type)", "양질토, 뻘흙, 암버럭, 혼합토 등 토사의 종류"),
        ("양질토 (GOOD_SOIL)", "이물질이 없고 성토용으로 가장 적합한 양질의 흙"),
        ("뻘흙 (MUD_SOIL)", "수분 함량이 너무 높아 질퍽거리는 반입 주의 토사"),
        ("암버럭 (ROCK)", "발파 및 천공 시 발생하는 크고 작은 돌덩어리와 바위더미"),
        ("혼합토 (MIXED)", "흙과 돌, 자갈이 무작위로 혼합된 상태의 토사"),
        ("오더 (Order / JobPost)", "기사를 모집하기 위해 등록하는 배차 요청 공고"),
        ("배차 (Dispatch)", "특정 오더에 차량과 기사를 매칭하여 운행을 지정하는 행위"),
        ("배차 티켓 (Dispatch Ticket)", "기사 1명의 1회 운행 주행 및 정산의 최소 시스템 단위"),
        ("콜 (Call)", "기사 앱에 실시간 수신되는 배차 수락 요청 알림"),
        ("미터기 (Fare Meter)", "실시간 GPS 속도와 시간을 병산해 덤프 운임을 산정하는 장치"),
        ("지오펜싱 (Geofencing)", "사토장 반경 200m 경계를 설정해 자동 도착을 감지하는 기술"),
        ("회차 (Return / Rejection)", "토질 불량 등으로 사토장에서 반입이 거부되어 되돌아가는 현상"),
        ("세륜기 (Washing Facility)", "비산먼지 방지를 위해 덤프 바퀴의 흙을 씻는 설비"),
        ("선등록 (Pre-registration)", "소장/차주가 번호만으로 사전에 기사/직원을 등록해 두는 기능"),
        ("매칭 (Matching)", "상차지 현장과 하차지 사토장을 상호 연결하는 중개 관계"),
        ("수용 공고 (Drop-Off Request)", "지주가 흙을 받겠다고 수량과 토종을 등록해두는 공고"),
        ("현장 관리자 (Site Manager / 소장)", "공사 현장을 개설하고 배차 오더를 생성 및 결제하는 책임자"),
        ("현장 담당자 (Site Worker / Staff)", "현장 게이트에서 덤프의 진출입을 승인하는 현장 실무 직원"),
        ("차주 (Owner / 사장님)", "덤프 차량을 보유하고 소속 기사 및 정산을 총괄하는 사업자"),
        ("기사 (Driver)", "덤프를 직접 운행하며 배차를 수행하고 미터기를 켜는 운전자"),
        ("하차지 지주 (Drop-Off Owner)", "사토장 땅을 소유하고 반입 토사 상태를 검사/승인하는 주체"),
        ("플랫폼 관리자 (Admin / 어드민)", "회원 심사, 단가/수수료 관리, 분쟁 발생 시 중재 정산 총괄"),
        ("15톤 덤프 (TRUCK_5T)", "소규모 도로 및 현장용 15톤 덤프 트럭 (약 10㎥ 적재)"),
        ("25톤 덤프 (TRUCK_25T)", "대형 토공사용 25.5톤 메인 덤프 트럭 (약 16~18㎥ 적재)"),
        ("앞사바리 (앞삼바리)", "앞바퀴 조향축이 2축인 대형 25톤 덤프 트럭의 현장 은어"),
        ("루베 (㎥)", "부피 단위인 세제곱미터. 현장에서 흙 양을 계산하는 단위"),
        ("공차", "적재함에 토사를 싣지 않고 운행하는 빈 트럭 상태"),
        ("덤프비 (운반비)", "운송의 대가로 기사/차주에게 지급하는 배차 운임 요금"),
        ("흙값 (토사 단가)", "운반비와 별개로 발생하며 흙 자체에 부과되는 원자재 대금"),
        ("처리비 (DROP_RECEIVE)", "뻘흙 등을 처리하는 사토장에 소장이 지주에게 지불하는 비용"),
        ("구입비 (DROP_PROVIDE)", "양질토를 얻기 위해 사토장 지주가 소장에게 지급하는 비용"),
        ("가상 지갑 (Wallet)", "기사/차주가 앱 내에서 운송료 정산액을 확인하고 출금하는 지갑"),
        ("직권 정산 (Forced Settlement)", "분쟁 발생 시 어드민이 GPS 경로를 확인해 강제 정산(70% 등)하는 조치"),
        ("공차 회차비", "반입 거부(회차)되어 돌아가는 기사에게 보존해 주는 최소 요금"),
        ("위약금 (Cancellation Penalty)", "매칭 완료 후 현장 사정으로 갑자기 취소 시 기사에게 주는 보상금")
    ]
    
    # Table creation
    table = doc.add_table(rows=1, cols=3)
    table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table)
    
    # Header row
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '현재명 (시스템 명칭)'
    hdr_cells[1].text = '바꿀 이름 (변경안 기재)'
    hdr_cells[2].text = '용어 설명 (Description)'
    
    for i, cell in enumerate(hdr_cells):
        set_cell_background(cell, 'D97706') # Amber Orange
        set_cell_margins(cell, top=140, bottom=140, left=150, right=150)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.bold = True
            run.font.name = 'Malgun Gothic'
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            
    col_widths = [Inches(2.5), Inches(2.2), Inches(3.3)]
            
    for idx, (curr_name, desc) in enumerate(terms):
        row = table.add_row()
        cells = row.cells
        cells[0].text = curr_name
        cells[1].text = "" 
        cells[2].text = desc
        
        bg_color = 'F9F8F6' if idx % 2 == 0 else 'FFFFFF'
        for i, cell in enumerate(cells):
            set_cell_background(cell, bg_color)
            set_cell_margins(cell, top=100, bottom=100, left=120, right=120)
            p = cell.paragraphs[0]
            for run in p.runs:
                run.font.name = 'Malgun Gothic'
                run.font.size = Pt(9)
                if i == 0:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(0x1F, 0x29, 0x37)
                else:
                    run.font.color.rgb = RGBColor(0x4B, 0x55, 0x63)
                    
    for row in table.rows:
        for i, width in enumerate(col_widths):
            row.cells[i].width = width

    # Save
    file_path = "D:\\Projects\\dumpring\\dumpring-platform-backend\\docs\\덤프링_핵심_용어_정의_및_변경_검토표.docx"
    doc.save(file_path)
    print(f"Glossary Word document successfully created at: {file_path}")

if __name__ == "__main__":
    main()
